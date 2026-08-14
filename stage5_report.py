"""Stage 5 — Auto-draft daily tracking report (Word).

Assembles the daily GTS tracking report: title page, key figures, team
deployment, time-spent analysis, coverage tables, statewide + per-LGA maps,
and placeholder sections for challenges & ERM photos that you fill in.
"""
import argparse
import os
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLUE = RGBColor(0x1F, 0x4E, 0x79)

# Usable width in inches for a picture on a portrait Letter page with the
# default 1in margins, and on the landscape section the maps get.
PORTRAIT_PICTURE_WIDTH = 6.3
LANDSCAPE_PICTURE_WIDTH = 9.2


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BLUE
    return h


def add_df_table(doc, df: pd.DataFrame, max_rows=30):
    df = df.head(max_rows)
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    for i, col in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(col)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            if isinstance(v, float):
                v = f"{v:,.1%}" if 0 <= v <= 1 and "%" in str(df.columns[i]) else f"{v:,.0f}"
            cells[i].text = "" if pd.isna(v) else str(v)
    return t


def _map_caption(filename: str, state: str, day: int | None = None) -> str:
    """Readable caption for a map PNG, from its filename.

    Filenames are `{State}_{LGA}_day_{N}.png` (LGA names go through
    `safe_name`, so '/' became '-' and spaces became underscores) and
    `{State}_statewide_day_{N}.png`.
    """
    stem = os.path.splitext(filename)[0]
    if state and stem.startswith(f"{state}_"):
        stem = stem[len(state) + 1:]
    if day is not None:
        stem = stem.replace(f"_day_{day}", "")
    name = stem.replace("_", " ").strip()
    day_suffix = f" — Day {day}" if day is not None else ""
    if name.lower() == "statewide":
        return f"{state} State — Settlement Visitation{day_suffix}"
    return f"{name} LGA{day_suffix}"


def _set_orientation(section, landscape: bool):
    """Point a section at landscape or portrait.

    python-docx does not swap the page dimensions when `orientation` is set, so
    both have to be assigned — setting only the orientation leaves a portrait
    text column on a page Word reports as landscape.
    """
    width, height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if (width > height) != landscape:
        section.page_width, section.page_height = height, width


def start_landscape_section(doc):
    """Begin a landscape section on a new page and return it.

    The side-by-side LGA maps are an A0 landscape page holding two panels; at
    the 6.3in a portrait page allows, each panel lands around 3in wide and the
    settlement names stop being legible. A landscape section gives them 9.2in,
    which is the difference between a map you can read and a thumbnail.
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_orientation(section, landscape=True)
    return section


def end_landscape_section(doc):
    """Return to portrait for whatever follows the maps."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_orientation(section, landscape=False)
    return section


def add_map_pages(doc, maps_folder: str, filenames, width: float = LANDSCAPE_PICTURE_WIDTH,
                  caption_of=None):
    """Embed a run of maps, one per page, inside a landscape section.

    Returns the number embedded. Does nothing (and starts no section) when
    there is nothing to show, so an empty maps folder cannot leave a stray
    blank landscape page in the report.
    """
    filenames = [f for f in filenames if os.path.exists(os.path.join(maps_folder, f))]
    if not filenames:
        return 0
    start_landscape_section(doc)
    for i, fn in enumerate(filenames):
        if i:
            doc.add_page_break()
        if caption_of:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(caption_of(fn))
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = BLUE
        doc.add_picture(os.path.join(maps_folder, fn), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_landscape_section(doc)
    return len(filenames)


def add_chart(doc, charts_folder, name, width=6.3):
    p = os.path.join(charts_folder, name) if charts_folder else None
    if p and os.path.exists(p):
        doc.add_picture(p, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def build_report(state: str, day: int, visitation_csv: str, cum_col: str,
                 deploy_csv: str | None, time_csv: str | None,
                 maps_folder: str, output_file: str, logo: str | None = None,
                 campaign_name: str = "Polio SIA", charts_folder: str | None = None,
                 state_logo: str | None = None) -> str:
    dip = pd.read_csv(visitation_csv, low_memory=False)
    lga_col = next(c for c in dip.columns if "lga" in c.lower() and "code" not in c.lower())

    visited = int((dip[cum_col] == "Visited").sum())
    not_visited = int((dip[cum_col] == "Not Visited").sum())
    planned = visited + not_visited
    pct = visited / planned * 100 if planned else 0

    doc = Document()

    # ---- title page
    if (logo and os.path.exists(logo)) or (state_logo and os.path.exists(state_logo)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if state_logo and os.path.exists(state_logo):
            p.add_run().add_picture(state_logo, width=Inches(1.2))
            if logo and os.path.exists(logo):
                p.add_run("     ")
        if logo and os.path.exists(logo):
            p.add_run().add_picture(logo, width=Inches(1.2))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{campaign_name}\nGIS Tracking of Vaccination Teams\n"
                      f"{state} State — Day {day} Daily Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(date.today().strftime("%d %B %Y")).font.size = Pt(14)
    doc.add_page_break()

    # ---- implementation map + background
    add_heading(doc, "Implementation Map")
    lga_n = dip[lga_col].astype(str).str.strip().str.title().nunique()
    ward_col = next((c for c in dip.columns if "ward" in c.lower() and "code" not in c.lower()), None)
    ward_n = int(dip.groupby(dip[lga_col].astype(str).str.strip().str.title())[ward_col]
                 .nunique().sum()) if ward_col else 0
    doc.add_paragraph(
        f"eHealth Africa is supporting the {state} campaign with vaccination tracking "
        f"across {lga_n} LGAs and {ward_n} wards. {planned:,} settlements planned "
        f"with geo-coordinates are being tracked.")
    impl_png = None
    if maps_folder and os.path.isdir(maps_folder):
        cand = [f for f in os.listdir(maps_folder) if "implementation" in f.lower()]
        if cand:
            impl_png = os.path.join(maps_folder, cand[0])
    if impl_png:
        doc.add_picture(impl_png, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- team deployment
    add_heading(doc, "Team Deployed / Team Reported")
    if deploy_csv and os.path.exists(deploy_csv):
        dep = pd.read_csv(deploy_csv)
        gt = dep[dep["LGA"] == "Grand Total"]
        if len(gt):
            g = gt.iloc[0]
            pct_txt = f" ({g['Reporting %']:.1%} reporting)" if "Reporting %" in g else ""
            doc.add_paragraph(
                f"{int(g['Teams Deployed']):,} teams were deployed and "
                f"{int(g['Teams Reported']):,} have reported as at the time of generating "
                f"this report, with {int(g['Teams Pending']):,} teams pending{pct_txt}.")
        add_chart(doc, charts_folder, "team_deployment.png", width=5.6)
        add_df_table(doc, dep)
    doc.add_paragraph("[Add narration on team deployment vs reporting compliance.]", style="Intense Quote")

    # ---- summary
    add_heading(doc, "Summary of Key Findings")
    doc.add_paragraph(
        f"On Day {day} of the campaign, GPS tracks were analysed against {planned:,} "
        f"planned settlements in {state} State. A total of {visited:,} settlements "
        f"({pct:.1f}%) showed evidence of vaccination team visitation, while "
        f"{not_visited:,} settlements ({100 - pct:.1f}%) were not yet visited.")
    doc.add_paragraph("[Add narration: notable improvements, LGAs of concern, follow-up actions agreed at the ERM.]",
                      style="Intense Quote")

    # ---- coverage charts
    add_heading(doc, f"State Daily Coverage — Day {day}")
    doc.add_paragraph(
        f"{visited:,} ({pct:.2f}%) of the {planned:,} settlements planned with "
        f"geo-coordinates were visited.")
    add_chart(doc, charts_folder, "state_coverage_donut.png", width=4.8)
    add_chart(doc, charts_folder, "lga_coverage_stacked.png", width=6.4)

    add_heading(doc, f"Cumulative Coverage by LGA — Day {day}")
    add_chart(doc, charts_folder, "lga_cumulative_counts.png", width=6.4)

    # ---- time spent
    add_heading(doc, "Time Spent Analysis (8:00am – 3:00pm)")
    add_chart(doc, charts_folder, "time_spent.png", width=5.8)
    if time_csv and os.path.exists(time_csv):
        add_df_table(doc, pd.read_csv(time_csv))
    doc.add_paragraph("[Add narration on field time compliance.]", style="Intense Quote")

    # ---- coverage per LGA
    add_heading(doc, "Settlement Visitation by LGA")
    visit = (dip.groupby([lga_col, cum_col]).size().unstack(fill_value=0))
    for c in ("Visited", "Not Visited"):
        if c not in visit.columns:
            visit[c] = 0
    visit = visit[["Visited", "Not Visited"]]
    visit["Total"] = visit.sum(axis=1)
    visit["% Visited"] = visit["Visited"] / visit["Total"]
    visit = visit.reset_index().rename(columns={lga_col: "LGA"})
    add_df_table(doc, visit)

    # ---- maps
    # These go in a landscape section, one per page: the per-LGA maps are wide
    # (two panels side by side when a side-by-side template is in use) and are
    # not readable squeezed into the portrait text column.
    add_heading(doc, "Visitation Coverage Maps")
    doc.add_paragraph(
        "Each LGA map shows settlement visitation status alongside gridded "
        "coverage for the same area.")
    if maps_folder and os.path.isdir(maps_folder):
        pngs = sorted(f for f in os.listdir(maps_folder)
                      if f.endswith(".png") and "implementation" not in f.lower())
        statewide = [f for f in pngs if "statewide" in f]
        lga_maps = [f for f in pngs if "statewide" not in f]
        add_map_pages(doc, maps_folder, statewide + lga_maps,
                      caption_of=lambda fn: _map_caption(fn, state, day))

    # ---- challenges & photos
    add_heading(doc, "Challenges and Recommendations")
    doc.add_paragraph("[List challenges observed today and recommendations.]", style="Intense Quote")
    add_heading(doc, "Deployment / Supportive Supervision Pictures")
    doc.add_paragraph("[Insert deployment, supervision and ERM photos here.]", style="Intense Quote")

    doc.save(output_file)
    print(f"Saved {output_file}")
    return output_file


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Draft daily tracking report")
    ap.add_argument("--state", required=True)
    ap.add_argument("--day", type=int, required=True)
    ap.add_argument("--visitation-csv", required=True)
    ap.add_argument("--cum-col", required=True)
    ap.add_argument("--deploy-csv", default=None)
    ap.add_argument("--time-csv", default=None)
    ap.add_argument("--maps-folder", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--logo", default=None)
    ap.add_argument("--campaign-name", default="Polio SIA")
    ap.add_argument("--state-logo", default=None, help="State PHC logo (title page)")
    a = ap.parse_args()
    build_report(a.state, a.day, a.visitation_csv, a.cum_col, a.deploy_csv,
                 a.time_csv, a.maps_folder, a.output, a.logo,
                 campaign_name=a.campaign_name, state_logo=a.state_logo)
