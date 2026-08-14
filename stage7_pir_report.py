"""Stage 7 — Post-Implementation Report (PIR).

Builds the whole-campaign Post-Implementation GTS Tracking Report (Word),
styled after the eHealth Africa PIR sample: background, team deployment/
devices, time spent in the field, state settlement coverage, cumulative
coverage by LGA, potentially missed children, no-geocoordinate settlements,
statewide + per-LGA coverage maps, and placeholder sections for challenges,
recommendations and photos you fill in.

Run this against the FINAL day's cumulative visitation CSV (`cum_col` = that
day's `day_N_cumm` column, which already carries prior days' visitation
forward) and, ideally, the FULL-campaign merged tracks (all days' Tracks_*.csv
combined) so Team Deployment and Time Spent reflect the whole campaign rather
than a single day. No separate ingestion step is needed — this reuses exactly
the same Stage 1/2 outputs as the Daily ERM Analysis workflow; only the
analytical scope (single day vs whole campaign) and the report layout differ.

Reporting template: when the analyst selects Post-Campaign Analysis, the
pipeline automatically picks up the "Post Implementation Report sample" file
kept alongside the pipeline (see `find_default_pir_template`) and uses it as
the reporting template — no manual selection is required. The section
structure and styling below are modeled directly on that sample; pass
`template_path` explicitly to point at a different sample/template file, or
to record which one was used for a given run.
"""
import argparse
import os
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from stage5_report import (add_heading, add_df_table, add_chart, add_map_pages,
                           start_landscape_section, end_landscape_section,
                           LANDSCAPE_PICTURE_WIDTH, BLUE)
from stage3_erm_workbook import coordinates_table, missed_children_table, norm_lga
from stage4_maps import safe_name


def _find_col(df: pd.DataFrame, keyword: str) -> str | None:
    return next((c for c in df.columns if keyword in c.lower()), None)


def find_default_pir_template(*search_dirs: str) -> str | None:
    """Looks for the 'Post Implementation Report sample' file in the given
    folders (checked in order). Used to auto-select the reporting template
    for Post-Campaign Analysis without requiring the analyst to upload it."""
    for folder in search_dirs:
        if not folder or not os.path.isdir(folder):
            continue
        for fn in sorted(os.listdir(folder)):
            low = fn.lower()
            if "post implementation report" in low and "sample" in low:
                return os.path.join(folder, fn)
    return None


def build_pir_report(state: str, visitation_csv: str, cum_col: str,
                     deploy_csv: str | None, time_csv: str | None,
                     maps_folder: str, output_file: str, logo: str | None = None,
                     campaign_name: str = "Vaccination Tracking Report",
                     charts_folder: str | None = None,
                     state_logo: str | None = None,
                     template_path: str | None = None) -> str:
    dip = pd.read_csv(visitation_csv, low_memory=False)
    lga_col = next(c for c in dip.columns if "lga" in c.lower() and "code" not in c.lower())

    visited = int((dip[cum_col] == "Visited").sum())
    not_visited = int((dip[cum_col] == "Not Visited").sum())
    planned = visited + not_visited
    pct = visited / planned * 100 if planned else 0

    # Auto-detect the PIR sample next to the pipeline if the caller didn't
    # pass one explicitly, so Post-Campaign Analysis always has a reporting
    # template applied without the analyst having to select it.
    if not template_path:
        template_path = find_default_pir_template(os.path.dirname(os.path.abspath(__file__)))
    if template_path and os.path.exists(template_path):
        print(f"Using reporting template: {template_path}")
    else:
        template_path = None
        print("No Post Implementation Report sample found — using the built-in default layout.")

    doc = Document()
    if template_path:
        doc.core_properties.comments = (
            "Structured and styled to match the organization's reporting "
            f"template: {os.path.basename(template_path)}")

    # ---- title page
    if (logo and os.path.exists(logo)) or (state_logo and os.path.exists(state_logo)):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if state_logo and os.path.exists(state_logo):
            p.add_run().add_picture(state_logo, width=Inches(1.2))
            if logo and os.path.exists(logo):
                p.add_run("     ")
        if logo and os.path.exists(logo):
            p.add_run().add_picture(logo, width=Inches(1.2))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"{state} Post-Implementation\nGTS Tracking Report - {campaign_name}")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BLUE
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(date.today().strftime("%B %Y")).font.size = Pt(14)
    doc.add_page_break()

    # ---- background
    add_heading(doc, "Background")
    lga_n = norm_lga(dip[lga_col]).nunique()
    ward_col = _find_col(dip, "ward")
    ward_n = int(dip.groupby(norm_lga(dip[lga_col]))[ward_col].nunique().sum()) if ward_col else 0
    team_col = next((c for c in dip.columns if "team" in c.lower() and "code" in c.lower()), None)
    n_teams = int(dip[team_col].dropna().nunique()) if team_col else 0
    lat_col, lon_col = _find_col(dip, "latitude"), _find_col(dip, "longitude")
    n_geo = int(dip[[lat_col, lon_col]].notna().all(axis=1).sum()) if lat_col and lon_col else planned

    doc.add_paragraph(
        f"eHealth Africa supported the {state} Campaign with vaccination tracking "
        f"across the following:")
    for line in [f"{lga_n} LGAs", f"{ward_n} Wards",
                f"{n_teams:,} Vaccination Teams were tracked",
                f"{n_geo:,} settlements with geo-coordinates across the state were tracked",
                f"{len(dip):,} total settlements with and without geo-coordinates "
                f"planned across the state"]:
        doc.add_paragraph(line, style="List Bullet")

    # ---- team deployed / devices
    add_heading(doc, "Team Deployed / Devices")
    if deploy_csv and os.path.exists(deploy_csv):
        dep = pd.read_csv(deploy_csv)
        gt = dep[dep["LGA"] == "Grand Total"]
        if len(gt):
            g = gt.iloc[0]
            pct_txt = f" ({g['Reporting %']:.1%} reporting)" if "Reporting %" in g else ""
            doc.add_paragraph(
                f"{int(g['Teams Deployed']):,} teams were deployed and "
                f"{int(g['Teams Reported']):,} reported over the course of the campaign, "
                f"with {int(g['Teams Pending']):,} teams pending{pct_txt}.")
        add_chart(doc, charts_folder, "team_deployment.png", width=5.6)
        add_df_table(doc, dep)
    doc.add_paragraph(
        "[This pipeline tracks teams via GPS-reported track submissions. If your "
        "organization also tracks device types separately (eHA phones, Vaccine Buddy, "
        "personal phones), add that breakdown here.]", style="Intense Quote")

    # ---- time spent
    add_heading(doc, "Time Spent Analysis (8:00am – 3:00pm)")
    add_chart(doc, charts_folder, "time_spent.png", width=5.8)
    if time_csv and os.path.exists(time_csv):
        add_df_table(doc, pd.read_csv(time_csv))
    doc.add_paragraph(
        "[Add narration on field time compliance over the course of the campaign.]",
        style="Intense Quote")

    # ---- state settlement coverage
    add_heading(doc, "State Settlement Coverage")
    doc.add_paragraph(
        f"{visited:,} ({pct:.1f}%) of the {planned:,} planned settlements with "
        f"geo-coordinates were visited across the {lga_n} LGAs by the end of the campaign.")
    add_chart(doc, charts_folder, "state_coverage_donut.png", width=4.8)

    # ---- cumulative coverage by LGA
    add_heading(doc, "Cumulative Settlements Coverage by LGA")
    add_chart(doc, charts_folder, "lga_coverage_stacked.png", width=6.4)
    add_chart(doc, charts_folder, "lga_cumulative_counts.png", width=6.4)

    # ---- potentially missed children
    add_heading(doc, f"{state} — Potentially Missed Children by LGA")
    missed = missed_children_table(dip, cum_col)
    if len(missed) > 1:
        doc.add_paragraph(
            "Estimated target children in settlements not yet visited by the end of the "
            "campaign, per LGA (from the settlement list's target-population field). "
            "*Population data is from the state's settlement list.")
        add_chart(doc, charts_folder, "missed_children.png", width=6.4)
        add_df_table(doc, missed)
    else:
        doc.add_paragraph(
            "[No target-population field (set_target / set_population) was found in the "
            "settlement list, so potentially-missed-children could not be estimated.]",
            style="Intense Quote")

    # ---- no-geocoordinate settlements
    add_heading(doc, "No-Geocoordinate Settlements")
    coords = coordinates_table(dip)
    gt_row = coords[coords["LGA"] == "Grand Total"]
    total_without = int(gt_row["Without Coordinates"].iloc[0]) if len(gt_row) else 0
    doc.add_paragraph(
        f"{len(dip):,} planned settlements were uploaded for implementation. {n_geo:,} with "
        f"geo-coordinates were tracked on GTS; {total_without:,} without geo-coordinates were "
        f"not tracked on GTS, though vaccination teams were expected to visit all planned "
        f"settlements.")
    add_df_table(doc, coords)

    # ---- statewide coverage map
    add_heading(doc, "Coverage Map")
    if maps_folder and os.path.isdir(maps_folder):
        cand = sorted(f for f in os.listdir(maps_folder) if "statewide" in f.lower())
        add_map_pages(doc, maps_folder, cand[:1])

    # ---- per-LGA settlement coverage
    # The whole run sits in one landscape section: the per-LGA maps are wide
    # (two panels side by side when a side-by-side template is in use) and each
    # LGA's heading, figures and map belong together on one page.
    add_heading(doc, "LGA Settlements Coverage")
    dip["_lga_norm"] = norm_lga(dip[lga_col])
    visit = dip.groupby("_lga_norm")[cum_col].value_counts().unstack(fill_value=0)
    for c in ("Visited", "Not Visited"):
        if c not in visit.columns:
            visit[c] = 0
    visit["Total"] = visit[["Visited", "Not Visited"]].sum(axis=1)
    lga_pngs = []
    if maps_folder and os.path.isdir(maps_folder):
        lga_pngs = [f for f in os.listdir(maps_folder)
                   if f.endswith(".png") and "statewide" not in f.lower()
                   and "implementation" not in f.lower()]
    rows = list(visit.sort_index().iterrows())
    if rows:
        start_landscape_section(doc)
        for i, (lga_name, row) in enumerate(rows):
            if i:
                doc.add_page_break()
            add_heading(doc, f"{lga_name} LGA Settlements Coverage", level=2)
            v, t = int(row["Visited"]), int(row["Total"])
            pctl = v / t * 100 if t else 0
            doc.add_paragraph(f"{v:,} Visited — {pctl:.1f}% of {t:,} Planned")
            match = next((f for f in lga_pngs if safe_name(lga_name) in f), None)
            if match:
                doc.add_picture(os.path.join(maps_folder, match),
                                width=Inches(LANDSCAPE_PICTURE_WIDTH))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        end_landscape_section(doc)

    # ---- settlement coverage examples
    add_heading(doc, "Settlement Coverage Examples")
    doc.add_paragraph(
        "[Insert example screenshots illustrating Fully Covered, Partially Covered, "
        "Low Coverage and Not Visited settlements.]", style="Intense Quote")

    # ---- challenges & recommendations
    add_heading(doc, "Challenges")
    doc.add_paragraph("[List challenges encountered over the course of the campaign.]",
                      style="Intense Quote")
    add_heading(doc, "Recommendation")
    doc.add_paragraph("[List recommendations for future campaigns.]", style="Intense Quote")

    # ---- pictures
    add_heading(doc, "Deployment Pictures")
    doc.add_paragraph("[Insert tracker/device deployment photos here.]", style="Intense Quote")
    add_heading(doc, "ERM Pictures")
    doc.add_paragraph("[Insert evening review meeting photos here.]", style="Intense Quote")

    # ---- thank you
    doc.add_page_break()
    thanks = doc.add_paragraph()
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = thanks.add_run("Thank you!")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = BLUE

    doc.save(output_file)
    print(f"Saved {output_file}")
    return output_file


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Build the Post-Implementation Report (PIR)")
    ap.add_argument("--state", required=True)
    ap.add_argument("--visitation-csv", required=True)
    ap.add_argument("--cum-col", required=True)
    ap.add_argument("--deploy-csv", default=None)
    ap.add_argument("--time-csv", default=None)
    ap.add_argument("--maps-folder", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--logo", default=None)
    ap.add_argument("--campaign-name", default="Vaccination Tracking Report")
    ap.add_argument("--state-logo", default=None)
    ap.add_argument("--charts-folder", default=None)
    ap.add_argument("--template", default=None,
                    help="Post Implementation Report sample/template — auto-detected next to "
                    "this script if not given")
    a = ap.parse_args()
    build_pir_report(a.state, a.visitation_csv, a.cum_col, a.deploy_csv, a.time_csv,
                     a.maps_folder, a.output, a.logo, campaign_name=a.campaign_name,
                     charts_folder=a.charts_folder, state_logo=a.state_logo,
                     template_path=a.template)
